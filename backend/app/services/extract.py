"""Extract plain text from PDF and DOCX — no paid APIs."""

from __future__ import annotations

import io
from pathlib import Path

from pypdf import PdfReader

from app.models.schema import MAX_EXTRACT_CHARS


def extract_pdf_text(file_bytes: bytes, max_chars: int = MAX_EXTRACT_CHARS) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        parts.append(t)
        if sum(len(p) for p in parts) >= max_chars:
            break
    text = "\n".join(parts).strip()
    return text[:max_chars]


def extract_docx_text(file_bytes: bytes, max_chars: int = MAX_EXTRACT_CHARS) -> str:
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    text = "\n".join(parts).strip()
    return text[:max_chars]


def extract_by_filename(filename: str, file_bytes: bytes, max_chars: int = MAX_EXTRACT_CHARS) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(file_bytes, max_chars)
    if suffix in (".docx",):
        return extract_docx_text(file_bytes, max_chars)
    raise ValueError(f"Unsupported file type: {suffix}")
